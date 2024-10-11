import tkinter as tk
from tkinter import scrolledtext, filedialog
import openai
import anthropic
import docx
import PyPDF2
import os
import pandas as pd
import json
import markdown
from bs4 import BeautifulSoup
from tkinter import messagebox
import re


class ChatApp:
    def __init__(self, master):
        self.master = master
        master.title("Chat Interface")

        self.client = openai.Client()
        self.claude_client = anthropic.Anthropic()
        self.messages = []

        self.document_uploaded = False
        self.document_text = ''
        self.document_texts = []
        self.uploaded_files = []

        # Dropdown for model selection
        self.model_var = tk.StringVar(master)
        self.model_var.set("gpt-4-turbo")  # default value
        self.model_options = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-4-0125-preview", "gpt-3.5-turbo", "gpt-3.5-turbo-1106", "claude-3-5-sonnet-20240620", "claude-3-opus-20240229", "claude-3-sonnet-20240229"]
        self.model_dropdown = tk.OptionMenu(master, self.model_var, *self.model_options)
        self.model_dropdown.grid(row=0, column=0, sticky="ew", padx=10, pady=10)

        self.which_model = self.model_var.get()

        self.model_message_map = {}  # 新增一个字典来存储消息ID和模型的映射关系

        self.current_processing_file = ''
        self.is_loading_previous_chat = False

        # Set the overall layout with grid
        master.grid_rowconfigure(0, weight=1)
        master.grid_columnconfigure(0, weight=1)

        # Scrolled Text box to show conversation
        self.conversation_box = scrolledtext.ScrolledText(master, width=200, height=32, wrap=tk.WORD)
        self.conversation_box.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky='nsew')

        # Using ScrolledText widget for user input, which includes built-in scrollbar functionality
        self.input_box = scrolledtext.ScrolledText(master, height=17, width=200, wrap=tk.WORD)
        self.input_box.grid(row=2, column=0, columnspan=2, padx=10, pady=10, sticky='nsew')

        # Button to send input
        self.send_button = tk.Button(master, text="Send", command=self.send_input)
        self.send_button.grid(row=3, column=1, padx=10, pady=5, sticky='e')

        # Frame for chat management buttons
        button_frame = tk.Frame(master)
        button_frame.grid(row=3, column=0, padx=10, pady=5, sticky='w')

        # Buttons for chat management
        self.start_new_chat_button = tk.Button(button_frame, text="Start New Chat", command=self.start_new_chat)
        self.start_new_chat_button.grid(row=0, column=0, padx=5, pady=5)

        self.save_button = tk.Button(button_frame, text="Save Chat", command=self.save_chat)
        self.save_button.grid(row=0, column=1, padx=5, pady=5)

        self.load_button = tk.Button(button_frame, text="Load Chat", command=self.load_chat)
        self.load_button.grid(row=0, column=2, padx=5, pady=5)

        self.upload_button = tk.Button(master, text="Upload Document", command=self.upload_document)
        self.upload_button.grid(row=4, column=0, columnspan=2, padx=10, pady=5, sticky='ew')

        self.file_list_box = tk.Listbox(master, height=3, width=60)
        self.file_list_box.grid(row=5, column=0, columnspan=2, padx=10, pady=5, sticky='ew')


    def send_input(self):
        user_input = self.input_box.get("1.0", "end-1c")  # Get input from Text widget
        conversation = self.conversation_box.get("1.0", "end-1c")
        if self.document_uploaded:
            user_input = user_input + "\n" + "File Content:" + self.document_text
            self.update_messages("user", user_input)
            self.input_box.delete("1.0", tk.END)
            self.invoke_api()
        elif user_input.strip():
            json_data = self.convert_to_json(conversation)
            self.messages = json_data["messages"]
            self.model_message_map = json_data["model_map"]
            self.update_messages("user", user_input)
            self.input_box.delete("1.0", tk.END)  # Clear the input box
            self.document_uploaded, self.document_text = False, ''
            self.invoke_api()
        elif conversation.strip() and not user_input.strip():
            json_data = self.convert_to_json(conversation)
            self.messages = json_data["messages"]
            self.model_message_map = json_data["model_map"]
            self.document_uploaded, self.document_text = False, ''
            self.invoke_api()



    def convert_to_json(self, input_string):
        messages = re.split(r'\n(?=(?:You|gpt-4-turbo|gpt-4o|gpt-4o-mini|gpt-4-0125-preview|gpt-3.5-turbo|claude-3-5-sonnet-20240620|claude-3-opus-20240229|claude-3-sonnet-20240229):)', input_string.strip())

        json_data = {"messages": [], "model_map": {}}
        assistant_models = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-4-0125-preview", "gpt-3.5-turbo",
                            "claude-3-5-sonnet-20240620", "claude-3-opus-20240229", "claude-3-sonnet-20240229"]

        for index, message in enumerate(messages):
            parts = message.split(': ', 1)
            role = parts[0].strip()
            content = parts[1].strip() if len(parts) > 1 else ""

            if role == "You":
                json_data["messages"].append({
                    "role": "user",
                    "content": [{"text": content, "type": "text"}]
                })
            elif role in assistant_models:
                json_data["messages"].append({
                    "role": "assistant",
                    "content": [{"text": content, "type": "text"}]
                })
                json_data["model_map"][str(index)] = role

        return json_data


    def invoke_api(self):
        self.which_model = self.model_var.get()
        if self.which_model in ["gpt-4o", "gpt-4-turbo", "gpt-4o-mini", "gpt-4-0125-preview", "gpt-3.5-turbo", "gpt-3.5-turbo-1106"]:
            try:
                response = self.client.chat.completions.create(
                    model=self.which_model,
                    messages=self.messages,
                    max_tokens=4095,
                )
                response_text = response.choices[0].message.content
                self.update_messages("assistant", response_text)
            except Exception as e:
                response_text =  str(e)
                self.update_messages("assistant", response_text)
        elif self.which_model in ["claude-3-5-sonnet-20240620", "claude-3-opus-20240229", "claude-3-sonnet-20240229"]:
            try:
                response = self.claude_client.messages.create(
                    model=self.which_model,
                    messages=self.messages,
                    max_tokens=4095,
                )
                response_text = response.content[0].text
                self.update_messages("assistant", response_text)
            except Exception as e:
                response_text =  str(e)
                self.update_messages("assistant", response_text)

        message_id = len(self.messages) - 1  # 使用消息数量-1作为最新消息的ID
        self.model_message_map[str(message_id)] = self.which_model


    def update_messages(self, role, text):
        self.messages.append({
            "role": role,
            "content": [{
                "text": text,
                "type": "text"
            }]
        })
        # Update conversation box
        who = "You: " if role == "user" else f"{self.which_model}: "
        self.conversation_box.insert(tk.END, who + text + '\n\n')
        self.conversation_box.yview(tk.END)


    def start_new_chat(self):
        self.messages = []
        self.is_loading_previous_chat = False
        # 初始化存储已上传文件名和内容的变量
        self.document_uploaded = False
        self.document_text = ''
        self.document_texts = []
        self.uploaded_files = []
        self.which_model = self.model_var.get()
        self.model_message_map = {}
        self.current_processing_file = ''

        self.input_box.delete("1.0", tk.END)
        self.conversation_box.delete('1.0', tk.END)
        self.file_list_box.delete(0, tk.END)


    def save_chat(self):
        if self.is_loading_previous_chat == False:
            default_save_path = "/"
            filename = filedialog.asksaveasfilename(
                initialdir=default_save_path,
                defaultextension=".json",
                filetypes=[("JSON files", "*.json")]
            )
        else:
            default_save_path = self.current_processing_file
            filename = filedialog.asksaveasfilename(
                initialfile=default_save_path,
                defaultextension=".json",
                filetypes=[("JSON files", "*.json")]
            )
        if filename:
            with open(filename, 'w') as file:
                json.dump({'messages': self.messages, 'model_map': self.model_message_map}, file)


    def load_chat(self):
        filename = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
        if filename:
            self.current_processing_file = filename
            self.is_loading_previous_chat = True
            with open(filename, 'r') as file:
                data = json.load(file)
                self.messages = data['messages']
                self.model_message_map = data['model_map']
                self.conversation_box.delete('1.0', tk.END)  

                for i, message in enumerate(self.messages):
                    role = message["role"]
                    text = message["content"][0]["text"]
                    if role == "assistant":
                        model_name = self.model_message_map[str(i)]
                        who = f"{model_name}: "
                    else:
                        who = "You: "
                    self.conversation_box.insert(tk.END, who + text + '\n\n')


    def upload_document(self):
        filenames = filedialog.askopenfilenames(
            title="Select files",
            filetypes=[("Documents", "*.docx *.pdf *.txt *.csv *.xlsx *.xls *.json *.md *.html")],
            initialdir="/",
        )
        if filenames:
            for filename in filenames:
                if len(self.uploaded_files) >= 4:
                    messagebox.showwarning("Limit Reached", "Maximum of 4 files can be uploaded.")
                    break
                if filename not in self.uploaded_files:
                    document_text = self.read_file(filename)
                    self.document_text = self.document_text + '\n\n' + f'File name: {filename}\nFile content:\n{document_text}'
                    self.uploaded_files.append(filename)
                    self.file_list_box.insert(tk.END, os.path.basename(filename))
                    self.document_uploaded = True
                    messagebox.showinfo("Success", f"{filename}has been successfully loaded!")



    def read_file(self, filename):
        if filename.endswith('.docx'):
            doc = docx.Document(filename)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            return '\n'.join(full_text)
        elif filename.endswith('.pdf'):
            with open(filename, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return ' '.join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
        elif filename.endswith('.txt'):
            with open(filename, 'r', encoding='utf-8') as file:
                return file.read()
        elif filename.endswith('.csv'):
            df = pd.read_csv(filename)
            return df.to_string(index=False)
        elif filename.endswith('.xlsx') or filename.endswith('.xls'):
            df = pd.read_excel(filename)
            return df.to_string(index=False)
        elif filename.endswith('.json'):
            with open(filename, 'r', encoding='utf-8') as file:
                data = json.load(file)
                return json.dumps(data, indent=2)
        elif filename.endswith('.md'):
            with open(filename, 'r', encoding='utf-8') as file:
                md_text = file.read()
                html = markdown.markdown(md_text)
                soup = BeautifulSoup(html, features="html.parser")
                return soup.get_text()
        elif filename.endswith('.html'):
            with open(filename, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file, 'html.parser')
                return soup.get_text()
        else:
            return ""


def main():
    root = tk.Tk()
    app = ChatApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
